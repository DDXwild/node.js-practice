import zipfile
from pathlib import Path
from shutil import copy2
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = Path(r"d:\Work\Node.js\Lab02\Node_Lab02.docx")
OUTPUT = ROOT / "Node_Lab03.docx"
REPO_URL = "https://github.com/JanRizhenko/node.js-practice"

STYLES = {
    "title": "Normal",
    "goal": "Text_lab Знак Знак Знак",
    "body": "ZIKS_MAIN_TEXT",
    "conclusion": "ZIKS_HEADER_FOURTH",
}

def normalize_dashes(text: str) -> str:
    return text.replace("\u2013", "-").replace("\u2014", "-").replace("\u2212", "-")


def clear_document_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=STYLES["title"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)


def add_goal(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=STYLES["goal"])
    label = p.add_run("Goal: ")
    label.bold = True
    p.add_run(text)


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=STYLES["body"])
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = True


def add_body(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph(style=STYLES["body"])
    run = p.add_run(text)
    run.bold = bold


def add_listing_label(doc: Document, label: str) -> None:
    p = doc.add_paragraph(style=STYLES["body"])
    run = p.add_run(label)
    run.bold = True


def add_code_lines(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style=STYLES["body"])
        run = p.add_run(line)
        run.font.size = Pt(8.5)


def add_file_listing(doc: Document, relative_path: str) -> None:
    file_path = ROOT / relative_path
    add_listing_label(doc, f"Listing {relative_path}:")
    content = normalize_dashes(file_path.read_text(encoding="utf-8"))
    add_code_lines(doc, content.splitlines())


def patch_header_xml(data: bytes) -> bytes:
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    root = ET.fromstring(data)
    text_nodes = list(root.iter(f"{{{ns}}}t"))

    for node in text_nodes:
        if node.text:
            node.text = normalize_dashes(node.text)
        if node.tail:
            node.tail = normalize_dashes(node.tail)

    for index, node in enumerate(text_nodes):
        if (node.text or "") == "\u041b\u0440" and index + 1 < len(text_nodes):
            next_node = text_nodes[index + 1]
            if (next_node.text or "") == "1":
                next_node.text = "3"

    xml_body = ET.tostring(root, encoding="unicode")
    return ("<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n" + xml_body).encode("utf-8")


def patch_template_headers(docx_path: Path) -> None:
    temp_path = docx_path.with_suffix(".tmp.docx")

    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename in ("word/header1.xml", "word/header2.xml"):
                    data = patch_header_xml(data)
                zout.writestr(item, data)

    temp_path.replace(docx_path)


def build_report() -> None:
    copy2(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)
    clear_document_body(doc)

    add_title(doc, "LABORATORY REPORT # 3")
    add_title(doc, "REST API with Express.js and integration testing")

    add_goal(
        doc,
        "The goal of this lab is to build a RESTful API on Express.js with TypeScript, "
        "validate input with Zod, use in-memory Map storage, split app.ts and server.ts, "
        "and write integration tests with Supertest.",
    )

    add_section_heading(doc, "Procedure:")

    add_body(doc, "Task 1. Project initialization", bold=True)
    add_body(
        doc,
        "Created the Lab03 project with express, cors, zod, typescript, jest, ts-jest, "
        "supertest, and npm scripts: dev, build, test, test:coverage.",
    )
    add_body(doc, "Project structure:")
    add_code_lines(
        doc,
        [
            "Lab03/",
            "  src/schemas/entity.schema.ts",
            "  src/storage/entity.ts",
            "  src/routes/entity.ts",
            "  src/middleware/validate.ts, errorHandler.ts",
            "  src/app.ts, server.ts",
            "  tests/entity.test.ts",
        ],
    )

    add_body(doc, "Task 2. Entity schemas and types (Movie)", bold=True)
    add_body(
        doc,
        "Entity: Movie. Fields: title (required), description (optional), releaseYear, "
        "genre (enum), director. updateSchema makes all fields optional. Entity type is "
        "inferred from createSchema and extended with id, createdAt, updatedAt.",
    )
    add_file_listing(doc, "src/schemas/entity.schema.ts")

    add_body(doc, "Task 3. In-memory storage", bold=True)
    add_body(
        doc,
        "storage/entity.ts uses Map<string, Entity> with getAll, getById, create, update, "
        "remove, reset, getRecent, and filtering at the storage layer.",
    )
    add_file_listing(doc, "src/storage/entity.ts")

    add_body(doc, "Task 4. CRUD routes and middleware", bold=True)
    add_body(
        doc,
        "validate(schema) parses req.body with Zod. errorHandler returns 400 for ZodError "
        "(err.issues) and 500 for other errors. Routes: GET/POST /api/movies, GET/PATCH/DELETE "
        "/api/movies/:id, GET /api/movies/recent.",
    )
    add_file_listing(doc, "src/middleware/validate.ts")
    add_file_listing(doc, "src/middleware/errorHandler.ts")
    add_file_listing(doc, "src/routes/entity.ts")

    add_body(doc, "Task 5. app.ts and server.ts", bold=True)
    add_body(doc, "app.ts configures Express without listen(). server.ts imports app and starts the server.")
    add_file_listing(doc, "src/app.ts")
    add_file_listing(doc, "src/server.ts")

    add_body(doc, "Task 6. Integration tests", bold=True)
    add_body(
        doc,
        "tests/entity.test.ts uses Supertest and beforeEach(storage.reset). "
        "15 tests cover CRUD, validation errors, 404 cases, query filters, and /recent.",
    )
    add_file_listing(doc, "tests/entity.test.ts")
    add_body(doc, "Test result: npm test - 15 passed. Code coverage: about 93% statements.")

    add_body(doc, "Task 7. Query filters and domain route", bold=True)
    add_body(
        doc,
        "Filters: genre, minYear, maxYear, title (combinable). Domain route GET /api/movies/recent "
        "returns movies released within the last 5 years.",
    )
    add_body(doc, "Example requests:")
    add_code_lines(
        doc,
        [
            "npm run dev",
            "curl -X POST http://localhost:3000/api/movies -H \"Content-Type: application/json\" \\",
            "  -d '{\"title\":\"Inception\",\"releaseYear\":2010,\"genre\":\"sci-fi\",\"director\":\"Christopher Nolan\"}'",
            "curl http://localhost:3000/api/movies?genre=sci-fi&minYear=2000",
            "curl http://localhost:3000/api/movies/recent",
        ],
    )

    add_body(doc, f"Repository link - {REPO_URL}")

    conclusion = doc.add_paragraph(style=STYLES["conclusion"])
    label = conclusion.add_run("Conclusion: ")
    label.bold = True
    conclusion.add_run(
        "A REST API for movies was implemented with correct HTTP status codes, Zod validation, "
        "in-memory storage, storage-level filtering, separated app/server setup, and 15 Supertest "
        "integration tests. The project is ready for submission."
    )

    doc.save(OUTPUT)
    patch_template_headers(OUTPUT)
    print(f"Report saved to {OUTPUT}")


if __name__ == "__main__":
    build_report()
